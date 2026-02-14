# src/tools/reporting_tools.py

import os
import smtplib
import logging
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import markdown
from fpdf import FPDF
import hashlib
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

from ..state import AgentState

# Helper to sanitize text (surrogate fix)
def sanitize_text(text):
    if not isinstance(text, str): return text
    # Encode to UTF-8 ignoring errors (strips surrogates), then decode back
    return text.encode('utf-8', 'replace').decode('utf-8')

# --------------------------------------------------------------------------
# NODO 3: GENERACIÓN DEL INFORME EN HTML
# --------------------------------------------------------------------------
def generate_report_node(state: AgentState) -> dict:
    """
    Toma los resúmenes y metadatos del estado para crear un informe completo en formato HTML.
    
    Args:
        state (AgentState): El estado actual del agente, que debe contener
                            'summaries', 'video_metadata' y 'topic'.

    Returns:
        dict: Un diccionario con la clave 'report' para actualizar el estado del agente.
    """
    logger.info("Generating report...")
    logger.info(f"State keys available in reporting: {list(state.keys())}")
    if "reddit_research" in state:
        logger.info(f"Reddit research results count: {len(state['reddit_research'])}")
    summaries = state.get("summaries", [])
    video_metadata = state.get("video_metadata", [])
    topic = state.get("original_topic", state.get("topic", "Tema desconocido"))
    consolidated = state.get("consolidated_summary", "")
    
    # Sanitize inputs immediately
    topic = sanitize_text(topic)
    consolidated = sanitize_text(consolidated)
    
    # Comprobar si tenemos CUALQUIER tipo de contenido para informar
    has_content = any([
        len(summaries) > 0,
        consolidated,
        state.get("wiki_research"),
        state.get("web_research"),
        state.get("arxiv_research"),
        state.get("scholar_research"),
        state.get("github_research"),
        state.get("hn_research"),
        state.get("so_research"),
        state.get("reddit_research")
    ])

    if not has_content:
        logger.warning("No information found to generate report.")
        report_html = f"<h1>Informe de Investigación sobre: {topic}</h1><p>No se encontró información relevante en las fuentes consultadas.</p>"
        return {"report": report_html}

    # Usamos f-strings de varias líneas para construir el HTML de manera legible.
    # Se añade un poco de estilo CSS en línea para mejorar la apariencia.
    # Premium HTML Design with modern CSS
    html_content = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Informe: {topic}</title>
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap" rel="stylesheet">
        <style>
            /* (Styles preserved) */
            * {{
                box-sizing: border-box;
            }}

            :root {{
                --primary: #2563eb;
                --primary-dark: #1e40af;
                --secondary: #64748b;
                --bg: #f8fafc;
                --card-bg: #ffffff;
                --text: #1e293b;
                --text-light: #64748b;
                --accent: #eff6ff;
                --border: #e2e8f0;
            }}
            
            body {{ 
                font-family: 'Inter', system-ui, -apple-system, sans-serif; 
                line-height: 1.6; 
                color: var(--text); 
                background-color: var(--bg);
                margin: 0;
                padding: 0;
            }}
            
            .container {{
                max-width: 900px;
                margin: 40px auto;
                padding: 0 20px;
            }}
            
            header {{
                text-align: center;
                margin-bottom: 50px;
                padding: 40px 0;
                background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
                color: white;
                border-radius: 16px;
                box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.1);
            }}
            
            h1 {{ margin: 0; font-weight: 700; font-size: 2.5rem; letter-spacing: -0.025em; }}
            .subtitle {{ opacity: 0.8; font-weight: 300; margin-top: 10px; font-size: 1.1rem; }}
            
            h2 {{ 
                color: var(--text); 
                font-size: 1.75rem; 
                margin-top: 40px; 
                margin-bottom: 20px;
                border-left: 4px solid var(--primary);
                padding-left: 15px;
            }}
            
            .section-card {{
                background: var(--card-bg);
                padding: 30px;
                border-radius: 12px;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                margin-bottom: 30px;
                border: 1px solid var(--border);
            }}
            
            .synthesis-card {{
                background: linear-gradient(to bottom right, #eff6ff, #ffffff);
                border: 1px solid #bfdbfe;
                padding: 40px;
            }}
            
            .synthesis-card h2 {{ border-left-color: #3b82f6; margin-top: 0; }}
            
            .research-item {{
                margin-bottom: 25px;
                padding-bottom: 20px;
                border-bottom: 1px solid var(--border);
            }}
            
            .research-item:last-child {{ border-bottom: none; margin-bottom: 0; padding-bottom: 0; }}
            
            .item-title {{ font-weight: 600; font-size: 1.25rem; color: var(--primary); margin-bottom: 8px; display: block; }}
            .item-meta {{ font-size: 0.875rem; color: var(--text-light); margin-bottom: 12px; }}
            .item-content {{ 
                font-size: 1rem; 
                color: var(--text); 
                word-break: break-word;
            }}
            
            a {{ color: var(--primary); text-decoration: none; font-weight: 500; }}
            a:hover {{ color: var(--primary-dark); text-decoration: underline; }}
            
            .tag {{
                display: inline-block;
                padding: 2px 10px;
                border-radius: 9999px;
                font-size: 0.75rem;
                font-weight: 600;
                background: var(--accent);
                color: var(--primary);
                margin-bottom: 10px;
            }}
            
            .summary-text {{ 
                word-break: break-word;
            }}
            .summary-text h2 {{
                font-size: 1.4rem;
                margin-top: 25px;
                margin-bottom: 12px;
                color: var(--primary-dark);
                border-left: none;
                padding-left: 0;
            }}
            .summary-text h3 {{
                font-size: 1.2rem;
                margin-top: 20px;
                margin-bottom: 10px;
                color: var(--primary);
                font-weight: 600;
            }}
            .summary-text ul {{
                padding-left: 25px;
                margin-bottom: 15px;
            }}
            .summary-text li {{
                margin-bottom: 8px;
            }}
            .summary-text p {{
                margin-bottom: 12px;
            }}
            
            .bib-list {{ list-style: none; padding: 0; }}
            .bib-list li {{ 
                padding: 12px 0; 
                border-bottom: 1px solid var(--border);
                font-size: 0.95rem;
                word-break: break-all;
            }}
            
            @media (max-width: 640px) {{
                h1 {{ font-size: 1.75rem; }}
                .container {{ margin: 20px auto; }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <header id="top">
                <h1>Investigación Inteligente</h1>
                <div class="subtitle">{topic}</div>
            </header>
    """

    # --- SECCIÓN: SÍNTESIS EJECUTIVA ---
    if state.get("consolidated_summary"):
        raw_summary = state["consolidated_summary"]
        
        # Pre-procesamiento para evitar listas numeradas planas: 
        # Convertimos cualquier línea que empiece por número (ej. "1. ") en una viñeta (*)
        processed_lines = []
        import re
        for line in raw_summary.split("\n"):
            # Si la línea empieza con un número y un punto, lo cambiamos por un asterisco
            # (dejando el resto de la línea igual, incluyendo posibles espacios de indentación)
            new_line = re.sub(r'^(\s*)\d+\.\s+', r'\1* ', line)
            processed_lines.append(new_line)
        
        processed_summary = "\n".join(processed_lines)
        synthesis_html = markdown.markdown(processed_summary)
        
        html_content += f"""
        <div class="section-card synthesis-card">
            <h2>💡 Síntesis Ejecutiva Consolidada</h2>
            <div class="summary-text">{synthesis_html}</div>
        </div>
        """

    # --- helper for safe strings ---
    def safe_str(val):
        return sanitize_text(str(val) if val is not None else "")

    # --- SECCIÓN: WIKIPEDIA ---
    if state.get("wiki_research"):
        html_content += "<h2 id='wiki'><span class='tag'>WIKIPEDIA</span> Contexto General</h2><div class='section-card'>"
        for item in state["wiki_research"]:
            summary = safe_str(item.get('summary', ''))
            if len(summary) > 500:
                summary = summary[:500] + "..."
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url'))}" class="item-title">{safe_str(item.get('title'))}</a>
                <p class="item-content">{summary}</p>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: WEB RESEARCH ---
    if state.get("web_research"):
        html_content += "<h2 id='web'><span class='tag'>WEB</span> Investigación Ampliada</h2><div class='section-card'>"
        for item in state["web_research"]:
            title = safe_str(item.get('title', 'Resultado Web'))
            content = safe_str(item.get('content', item.get('snippet', '')))
            if len(content) > 500:
                content = content[:500] + "..."
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url'))}" class="item-title">{title}</a>
                <p class="item-content">{content}</p>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: ARXIV ---
    if state.get("arxiv_research"):
        html_content += "<h2 id='arxiv'><span class='tag'>ACADÉMICO</span> Artículos en arXiv</h2><div class='section-card'>"
        for item in state["arxiv_research"]:
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url', '#'))}" class="item-title">{safe_str(item.get('title'))}</a>
                <div class="item-meta">Autores: {safe_str(item.get('authors'))}</div>
                <p class="item-content">{safe_str(item.get('summary'))}</p>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: SEMANTIC SCHOLAR ---
    if state.get("scholar_research"):
        html_content += "<h2 id='scholar'><span class='tag'>CIENCIA</span> Semantic Scholar</h2><div class='section-card'>"
        for item in state["scholar_research"]:
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url', '#'))}" class="item-title">{safe_str(item.get('title'))} ({safe_str(item.get('year', 'N/A'))})</a>
                <div class="item-meta">Autores: {safe_str(item.get('authors'))}</div>
                <p class="item-content">{safe_str(item.get('content'))}</p>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: GITHUB ---
    if state.get("github_research"):
        html_content += "<h2 id='github'><span class='tag'>CÓDIGO</span> Repositorios Destacados</h2><div class='section-card'>"
        for item in state["github_research"]:
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url'))}" class="item-title">{safe_str(item.get('name'))} (⭐ {safe_str(item.get('stars'))})</a>
                <p class="item-content">{safe_str(item.get('description'))}</p>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: HACKER NEWS ---
    if state.get("hn_research"):
        html_content += "<h2 id='hn'><span class='tag'>HACKER NEWS</span> Discusiones</h2><div class='section-card'>"
        for item in state["hn_research"]:
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url'))}" class="item-title">{safe_str(item.get('title'))}</a>
                <div class="item-meta">Autor: {safe_str(item.get('author'))} | Puntos: {safe_str(item.get('points'))}</div>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: STACK OVERFLOW ---
    if state.get("so_research"):
        html_content += "<h2 id='so'><span class='tag'>STACK OVERFLOW</span> Soporte Técnico</h2><div class='section-card'>"
        for item in state["so_research"]:
            html_content += f"""
            <div class="research-item">
                <a href="{safe_str(item.get('url'))}" class="item-title">{safe_str(item.get('title'))}</a>
                <div class="item-meta">Score: {safe_str(item.get('score'))} | Resuelta: {'Sí' if item.get('is_answered') else 'No'}</div>
                <div class="tag-container">
                    {' '.join([f'<span class="tag">{safe_str(t).strip()}</span>' for t in str(item.get('tags', '')).split(',')])}
                </div>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"
    
    # --- SECCIÓN: REDDIT ---
    if state.get("reddit_research"):
        html_content += "<h2 id='reddit'><span class='tag'>REDDIT</span> Discusiones y Opiniones</h2><div class='section-card'>"
        for item in state["reddit_research"]:
            content = safe_str(item.get('content', item.get('snippet', '')))
            if len(content) > 500:
                content = content[:500] + "..."
            html_content += f"""
            <div class="research-item">
                <p class="item-content">{content}</p>
                <a href="{safe_str(item.get('url'))}" class="item-meta">Ver hilo en Reddit &rarr;</a>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: YOUTUBE ---
    if state.get("local_research"):
        html_content += "<h2 id='local'><span class='tag'>LOCAL</span> Conocimiento Interno</h2><div class='section-card'>"
        for item in state["local_research"]:
            content = safe_str(item.get('content', ''))
            if len(content) > 500:
                content = content[:500] + "..."
            html_content += f"""
            <div class="research-item">
                <div class="item-title">{safe_str(item.get('title'))}</div>
                <p class="item-content">{content}</p>
                <a href="{safe_str(item.get('url'))}" class="item-meta">Ver archivo local &rarr;</a>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    if summaries:
        html_content += "<h2 id='youtube'><span class='tag'>MULTIMEDIA</span> Análisis de YouTube</h2><div class='section-card'>"
        for i, (summary, metadata) in enumerate(zip(summaries, video_metadata)):
            html_content += f"""
            <div class="research-item">
                <div class="item-title">Vídeo {i+1}: {safe_str(metadata.get('title', 'Sin título'))}</div>
                <div class="item-meta">Autor: {safe_str(metadata.get('author', 'Desconocido'))} | <a href="{safe_str(metadata.get('url', '#'))}">Ver en YouTube</a></div>
                <div class="item-content summary-text">{safe_str(summary)}</div>
            </div>
            """
        html_content += "<a href='#top' style='font-size: 0.8rem;'>&uarr; Volver al inicio</a></div>"

    # --- SECCIÓN: BIBLIOGRAFÍA ---
    bibliography = []
    has_bib = any([
        state.get("wiki_research"),
        state.get("arxiv_research"),
        state.get("scholar_research"),
        state.get("github_research"),
        state.get("hn_research"),
        state.get("so_research"),
        state.get("reddit_research"),
        state.get("local_research"),
        video_metadata
    ])
    
    if has_bib:
        html_content += "<hr><h2>📚 Bibliografía y Fuentes</h2><div class='section-card'><ul class='bib-list'>"
        
        # Wiki
        for item in state.get("wiki_research", []):
            url = safe_str(item.get('url', '#'))
            title = safe_str(item.get('title', 'Wikipedia'))
            ref = f"Wikipedia: {title} - {url}"
            bibliography.append(ref)
            html_content += f"<li>Wikipedia: {title} - <a href='{url}'>{url}</a></li>"
        # arXiv
        for item in state.get("arxiv_research", []):
            url = safe_str(item.get('url', '#'))
            title = safe_str(item.get('title', 'Articulo arXiv'))
            authors = safe_str(item.get('authors', 'Desconocido'))
            ref = f"arXiv: {title} ({authors}) - {url}"
            bibliography.append(ref)
            html_content += f"<li>arXiv: {title} ({authors}) - <a href='{url}'>{url}</a></li>"
        # Scholar
        for item in state.get("scholar_research", []):
            url = safe_str(item.get('url', '#'))
            title = safe_str(item.get('title', 'Articulo Scholar'))
            year = safe_str(item.get('year', 'N/A'))
            ref = f"Semantic Scholar: {title} ({year}) - {url}"
            bibliography.append(ref)
            html_content += f"<li>Semantic Scholar: {title} ({year}) - <a href='{url}'>{url}</a></li>"
        # GitHub
        for item in state.get("github_research", []):
            url = safe_str(item.get('url', '#'))
            name = safe_str(item.get('name', 'Repository'))
            ref = f"GitHub: {name} - {url}"
            bibliography.append(ref)
            html_content += f"<li>GitHub: {name} - <a href='{url}'>{url}</a></li>"
        # Hacker News
        for item in state.get("hn_research", []):
            url = safe_str(item.get('url', '#'))
            title = safe_str(item.get('title', 'Hacker News'))
            ref = f"Hacker News: {title} - {url}"
            bibliography.append(ref)
            html_content += f"<li>Hacker News: {title} - <a href='{url}'>{url}</a></li>"
        # Stack Overflow
        for item in state.get("so_research", []):
            url = safe_str(item.get('url', '#'))
            title = safe_str(item.get('title', 'Stack Overflow'))
            ref = f"Stack Overflow: {title} - {url}"
            bibliography.append(ref)
            html_content += f"<li>Stack Overflow: {title} - <a href='{url}'>{url}</a></li>"
        # Reddit
        for item in state.get("reddit_research", []):
            url = safe_str(item.get('url', '#'))
            title = "Discusion en Reddit"
            ref = f"Reddit: {title} - {url}"
            bibliography.append(ref)
            html_content += f"<li>Reddit: {title} - <a href='{url}'>{url}</a></li>"
        # YouTube
        for metadata in video_metadata:
            url = safe_str(metadata.get('url', '#'))
            title = safe_str(metadata.get('title', 'Video'))
            author = safe_str(metadata.get('author', 'Autor'))
            ref = f"YouTube: {title} por {author} - {url}"
            bibliography.append(ref)
            html_content += f"<li>YouTube: {title} por {author} - <a href='{url}'>{url}</a></li>"
        
        # Local knowledge
        for item in state.get("local_research", []):
            url = safe_str(item.get('url', '#'))
            title = safe_str(item.get('title', 'Archivo Local'))
            ref = f"Local: {title} - {url}"
            bibliography.append(ref)
            html_content += f"<li>Local: {title} - <a href='{url}'>{url}</a></li>"
        
        html_content += "</ul></div>"

    html_content += """
    </div>
    </body>
    </html>
    """
    
    # Preparamos los textos para otros formatos
    markdown_text = f"# Informe de Investigación: {topic}\n\n"
    if consolidated:
        markdown_text += f"## Síntesis Ejecutiva\n{consolidated}\n\n"
    
    markdown_text += "## Bibliografía\n"
    for ref in bibliography:
        markdown_text += f"- {ref}\n"
    
    # Sanitize markdown
    markdown_text = sanitize_text(markdown_text)

    # Ensure reports directory exists
    reports_dir = "reports"
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)

    # Guardamos el HTML
    report_path = os.path.join(reports_dir, "reporte_final.html")
    # ENFORCE SANITIZATION ON FINAL WRITE
    try:
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(sanitize_text(html_content))
    except Exception as e:
        logger.error(f"Error saving HTML (surrogate check): {e}")
    
    # Guardamos el Markdown
    # Sanitize topic for filename (prevent path injection / Errno 2)
    safe_topic = topic.replace(" ", "_").replace("/", "_").replace("\\", "_")[:30]
    md_path = os.path.join(reports_dir, f"reporte_{safe_topic}.md")
    try:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
    except Exception as e:
        logger.error(f"Error saving MD (surrogate check): {e}")
    
    # Generamos el DOCX
    docx_path = os.path.join(reports_dir, "reporte_final.docx")
    try:
        generate_docx(state, topic, docx_path, bibliography)
        logger.info("DOCX generated successfully.")
    except Exception as e:
        logger.warning("docx_generation_failed", exc_info=e)
        docx_path = None

    # --- GENERACIÓN DE PDF ---
    pdf_path = os.path.join(reports_dir, "reporte_investigacion.pdf")
    try:
        generate_pdf(state, topic, pdf_path, bibliography) 
        logger.info("pdf_generated")
    except Exception as e:
        logger.warning("pdf_generation_failed", exc_info=e)
        pdf_path = None

    # Phase 6: Save finalized research state to database for persistence
    try:
        from db_manager import save_session
        save_session(topic, state.get("persona", "general"), state)
        logger.info(f"✅ Research session persisted for: {topic}")
    except Exception as e_db:
        logger.error(f"Failed to persist research session: {e_db}")

    logger.info("html_report_generated")
    return {
        "report": html_content, 
        "bibliography": bibliography, 
        "pdf_path": pdf_path,
        "md_path": md_path,
        "docx_path": docx_path
    }

def generate_docx(state: AgentState, topic: str, output_path: str, bibliography: list):
    """Genera un archivo Word (.docx) profesional."""
    doc = Document()
    doc.add_heading(f'Informe de Investigación: {topic}', 0)
    
    if state.get("consolidated_summary"):
        doc.add_heading('Síntesis Ejecutiva', level=1)
        doc.add_paragraph(state["consolidated_summary"])
        
    doc.add_heading('Bibliografía', level=1)
    for ref in bibliography:
        doc.add_paragraph(ref, style='List Bullet')
        
    doc.save(output_path)

def generate_pdf(state: AgentState, topic: str, output_path: str, bibliography_list: list = None):
    """Genera un archivo PDF profesional usando fpdf2 con todas las secciones."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Margen explícito
    l_margin = 15
    pdf.set_left_margin(l_margin)
    pdf.set_right_margin(l_margin)
    eff_w = pdf.w - 2 * l_margin

    # Usamos Helvetica (estándar). No admite emojis ni caracteres especiales complejos.
    pdf.set_font("Helvetica", "B", 16)
    
    # Título
    safe_topic = topic.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(eff_w, 10, f"Informe de Investigacion: {safe_topic}", align='C')
    pdf.ln(5)
    
    def clean_text(text):
        if not text: return ""
        # Removemos acentos problemáticos y emojis
        import unicodedata
        text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
        return text

    def add_section_header(title):
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.multi_cell(eff_w, 10, clean_text(title), fill=True)
        pdf.ln(2)

    # Síntesis
    add_section_header("Sintesis Ejecutiva Consolidada")
    
    summary_md = state.get("consolidated_summary", "No disponible")
    
    # Simple Markdown Parser for PDF
    for line in summary_md.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(2)
            continue
            
        # Limpieza básica de negritas Markdown (**) para el PDF simple
        line = line.replace("**", "")

        if line.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(eff_w, 7, clean_text(line.replace("### ", "")))
            pdf.ln(1)
        elif line.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(eff_w, 8, clean_text(line.replace("## ", "")))
            pdf.ln(1)
        elif line.startswith("* ") or line.startswith("- "):
            pdf.set_font("Helvetica", "", 10)
            # Indentación para viñetas
            pdf.set_x(l_margin + 7)
            # Extraemos el contenido quitando el símbolo de viñeta
            content = line[2:].strip()
            pdf.multi_cell(eff_w - 7, 6, f"- {clean_text(content)}")
        elif any(line.startswith(f"{i}. ") for i in range(1, 20)):
            # Fallback para listas numeradas si el LLM las genera
            pdf.set_font("Helvetica", "", 10)
            pdf.set_x(l_margin + 7)
            pdf.multi_cell(eff_w - 7, 6, clean_text(line))
        else:
            # Texto normal o subtemas que vienen formateados como "1. Titulo"
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(eff_w, 6, clean_text(line))
    
    pdf.ln(5)


    pdf.output(output_path)


# --------------------------------------------------------------------------
# NODO 4: ENVÍO DEL INFORME POR CORREO ELECTRÓNICO
# --------------------------------------------------------------------------
def send_email_node(state: AgentState) -> dict:
    """
    Envía el informe generado por correo electrónico utilizando las credenciales del archivo .env.
    Implementa idempotencia para evitar duplicados en la misma sesión/tema.

    Args:
        state (AgentState): El estado actual del agente, que contiene el 'report' en HTML.

    Returns:
        dict: Un diccionario con una bandera 'email_sent' para evitar re-envíos.
    """
    logger.info("send_email_node_started")
    
    # Verificación de idempotencia
    report = state.get("report", "")
    topic = state.get("original_topic", state.get("topic", "Informe-Investigacion"))
    file_topic = topic.replace(" ", "_")[:50]
    if not report:
        logger.warning("no_report_to_send")
        return {}
        
    # Crear un hash del reporte para identificar envíos duplicados
    report_hash = hashlib.md5(report.encode('utf-8')).hexdigest()
    
    # Comprobar si ya enviamos este reporte exacto en esta ejecución
    if state.get("last_email_hash") == report_hash:
        logger.info("email_already_sent", report_hash=report_hash)
        return {}

    # Obtenemos la configuración del correo desde las variables de entorno.
    from ..config import settings
    sender_email = settings.email_username
    receiver_email = settings.email_recipient
    password = settings.email_password
    host = settings.email_host
    port = settings.email_port         

    if not all([sender_email, receiver_email, password]):
        logger.error("email_credentials_missing")
        return {}

    # Creación del objeto del mensaje de correo.
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = f"Informe de Investigación: {topic}"
    msg.attach(MIMEText(report, 'html'))

    # Adjuntamos el PDF si existe
    pdf_path = state.get("pdf_path")
    if pdf_path and os.path.exists(pdf_path):
        try:
            with open(pdf_path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())
            
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename={os.path.basename(pdf_path)}",
            )
            msg.attach(part)
            logger.info("pdf_attached_to_email", path=pdf_path)
        except Exception as e:
            logger.warning("pdf_attachment_failed", exc_info=e)

    try:
        # Iniciamos la conexión con el servidor SMTP.
        logger.info("smtp_connecting", host=host, port=port)
        server = smtplib.SMTP(host, port, timeout=30)
        server.starttls()  
        server.login(sender_email, password)
        
        # Enviamos el correo.
        server.sendmail(sender_email, receiver_email, msg.as_string())
        logger.info("email_sent", recipient=receiver_email)
        
        # Devolvemos el hash para evitar envíos futuros del mismo contenido
        return {"last_email_hash": report_hash}
        
    except smtplib.SMTPAuthenticationError:
        logger.error("smtp_authentication_failed")
        return {}
    except Exception as e:
        logger.error("email_send_failed", exc_info=e)
        return {}
    finally:
        if 'server' in locals() and server.sock:
            server.quit() 